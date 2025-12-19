import modal
from datetime import datetime

app = modal.App("selatek-notes-multiuser")

image = modal.Image.debian_slim().pip_install("fastapi", "python-docx", "Pillow")

# Admin-lösenord för att skapa nya användare
ADMIN_SECRET = "selatek2024"

# Modal Dict för användare och sessioner
users = modal.Dict.from_name("selatek-users", create_if_missing=True)
sessions = modal.Dict.from_name("selatek-sessions", create_if_missing=True)


def get_user_html(username, has_pin):
    """Generera HTML för en specifik användare"""
    has_pin_js = 'true' if has_pin else 'false'
    return f'''<!DOCTYPE html>
<html lang="sv">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, user-scalable=no">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="apple-mobile-web-app-title" content="Selatek - {username}">
    <title>Selatek Notes - {username}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            min-height: 100vh;
            color: #fff;
            padding: 20px;
            padding-top: env(safe-area-inset-top, 20px);
            padding-bottom: env(safe-area-inset-bottom, 20px);
        }}
        .container {{ max-width: 600px; margin: 0 auto; }}
        h1 {{ text-align: center; margin-bottom: 5px; font-size: 24px; }}
        .user-badge {{ text-align: center; font-size: 14px; color: #60a5fa; margin-bottom: 15px; }}
        .status {{ text-align: center; font-size: 12px; color: #888; margin-bottom: 20px; }}
        .status.online {{ color: #4ade80; }}
        .status.offline {{ color: #f87171; }}
        .text-area {{
            width: 100%; min-height: 250px;
            background: rgba(255,255,255,0.1);
            border: 1px solid rgba(255,255,255,0.2);
            border-radius: 12px; padding: 15px;
            color: #fff; font-size: 16px; line-height: 1.6;
            resize: none; margin-bottom: 15px;
        }}
        .text-area:focus {{ outline: none; border-color: #60a5fa; }}
        .text-area::placeholder {{ color: rgba(255,255,255,0.4); }}
        .image-preview {{
            margin-bottom: 15px;
        }}
        .image-counter-big {{
            display: inline-block;
            background: #fff;
            color: #000;
            padding: 8px 16px;
            border-radius: 8px;
            font-size: 16px;
            font-weight: bold;
        }}
        .buttons {{ display: flex; flex-direction: column; gap: 10px; }}
        .btn-row {{ display: flex; gap: 10px; }}
        .btn-row button {{ flex: 1; }}
        button {{
            width: 100%; padding: 16px; border: none; border-radius: 12px;
            font-size: 16px; font-weight: 600; cursor: pointer;
            transition: transform 0.1s, opacity 0.2s;
        }}
        button:active {{ transform: scale(0.98); }}
        button:disabled {{ opacity: 0.5; cursor: not-allowed; }}
        .btn-dictate {{ background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); color: #fff; }}
        .btn-dictate.recording {{
            background: linear-gradient(135deg, #ef4444 0%, #b91c1c 100%);
            animation: pulse 1s infinite;
        }}
        @keyframes pulse {{ 0%, 100% {{ opacity: 1; }} 50% {{ opacity: 0.7; }} }}
        .btn-photo {{ background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); color: #fff; }}
        .btn-save {{ background: linear-gradient(135deg, #10b981 0%, #047857 100%); color: #fff; }}
        .btn-clear {{ background: rgba(255,255,255,0.1); color: #fff; border: 1px solid rgba(255,255,255,0.2); }}
        .btn-pin {{ background: linear-gradient(135deg, #8b5cf6 0%, #6d28d9 100%); color: #fff; margin-top: 20px; }}
        .word-count {{ text-align: center; font-size: 12px; color: #888; margin-top: 10px; }}
        .toast {{
            position: fixed; bottom: 100px; left: 50%; transform: translateX(-50%);
            background: rgba(0,0,0,0.8); color: #fff; padding: 12px 24px;
            border-radius: 8px; font-size: 14px; opacity: 0;
            transition: opacity 0.3s; z-index: 1000;
        }}
        .toast.show {{ opacity: 1; }}
        #cameraInput {{ display: none; }}

        /* PIN-modal */
        .pin-overlay {{
            position: fixed; top: 0; left: 0; right: 0; bottom: 0;
            background: rgba(0,0,0,0.9); z-index: 2000;
            display: flex; align-items: center; justify-content: center;
        }}
        .pin-overlay.hidden {{ display: none; }}
        .pin-box {{
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            padding: 30px; border-radius: 16px; text-align: center;
            max-width: 300px; width: 90%;
        }}
        .pin-box h2 {{ margin-bottom: 20px; }}
        .pin-box input {{
            width: 100%; padding: 15px; font-size: 24px; text-align: center;
            border: 2px solid rgba(255,255,255,0.2); border-radius: 12px;
            background: rgba(255,255,255,0.1); color: #fff;
            margin-bottom: 15px; letter-spacing: 8px;
        }}
        .pin-box input:focus {{ outline: none; border-color: #60a5fa; }}
        .pin-box .pin-error {{ color: #f87171; font-size: 14px; margin-bottom: 10px; }}
        .pin-box button {{
            background: linear-gradient(135deg, #10b981 0%, #047857 100%);
        }}
        .pin-info {{ font-size: 12px; color: #888; margin-top: 10px; }}
    </style>
</head>
<body>
    <!-- PIN-modal -->
    <div class="pin-overlay" id="pinOverlay">
        <div class="pin-box">
            <h2 id="pinTitle">Ange PIN-kod</h2>
            <div class="pin-error" id="pinError"></div>
            <input type="password" id="pinInput" maxlength="4" pattern="[0-9]*" inputmode="numeric" placeholder="••••">
            <button onclick="submitPin()">OK</button>
            <div class="pin-info" id="pinInfo"></div>
        </div>
    </div>

    <div class="container" id="mainApp">
        <h1>Selatek Notes</h1>
        <div class="user-badge">Inloggad som: {username}</div>
        <div class="status" id="status">Kontrollerar...</div>
        <textarea class="text-area" id="textArea" placeholder="Tryck på mikrofonen för att diktera...">Objekt: </textarea>
        <div class="image-preview" id="imagePreview"></div>
        <div class="buttons">
            <div class="btn-row">
                <button class="btn-dictate" id="dictateBtn">🎤 Diktera</button>
                <button class="btn-photo" id="photoBtn">📷 Foto</button>
            </div>
            <button class="btn-save" id="saveBtn">💾 Spara till Word</button>
            <button class="btn-clear" id="clearBtn">🗑️ Rensa</button>
            <button class="btn-pin" id="pinBtn">🔐 Ändra PIN-kod</button>
        </div>
        <div class="word-count" id="wordCount">0 tecken | 0 bilder</div>
    </div>
    <input type="file" id="cameraInput" accept="image/*" capture="environment">
    <div class="toast" id="toast"></div>
    <script>
        const HAS_SERVER_PIN = {has_pin_js};
        const USERNAME = '{username}';
        const ACCUMULATE_URL = 'https://mackanh1972--selatek-notes-multiuser-accumulate.modal.run';
        const IMAGE_URL = 'https://mackanh1972--selatek-notes-multiuser-add-image.modal.run';
        const WORD_URL = 'https://mackanh1972--selatek-notes-multiuser-save-to-word.modal.run';
        const STORAGE_KEY = 'selatek-' + USERNAME + '-text';
        const IMAGES_KEY = 'selatek-' + USERNAME + '-images';
        const SESSION_KEY = 'selatek-' + USERNAME + '-session';

        const textArea = document.getElementById('textArea');
        const dictateBtn = document.getElementById('dictateBtn');
        const photoBtn = document.getElementById('photoBtn');
        const saveBtn = document.getElementById('saveBtn');
        const clearBtn = document.getElementById('clearBtn');
        const statusEl = document.getElementById('status');
        const wordCountEl = document.getElementById('wordCount');
        const toast = document.getElementById('toast');
        const cameraInput = document.getElementById('cameraInput');
        const imagePreview = document.getElementById('imagePreview');

        let isRecording = false;
        let recognition = null;

        function getSessionId() {{
            // Använd alltid samma session per användare (inte slumpmässig)
            // Detta gör att bilder finns kvar även om localStorage rensas
            return 'default';
        }}

        function saveLocal() {{
            try {{
                localStorage.setItem(STORAGE_KEY, textArea.value);
            }} catch (e) {{
                console.error('localStorage error:', e);
            }}
            updateWordCount();
        }}

        function loadLocal() {{
            const saved = localStorage.getItem(STORAGE_KEY);
            if (saved) {{
                textArea.value = saved;
            }} else {{
                textArea.value = 'Objekt: ';
            }}
            updateWordCount();
        }}

        function updateStatus() {{
            statusEl.textContent = navigator.onLine ? '🟢 Online' : '🔴 Offline';
            statusEl.className = 'status ' + (navigator.onLine ? 'online' : 'offline');
        }}

        function updateWordCount() {{
            wordCountEl.textContent = textArea.value.length + ' tecken | ' + imageCount + ' bilder';
        }}

        function showToast(msg) {{
            toast.textContent = msg;
            toast.classList.add('show');
            setTimeout(() => toast.classList.remove('show'), 3000);
        }}


        function setupSpeechRecognition() {{
            if (!('webkitSpeechRecognition' in window) && !('SpeechRecognition' in window)) {{
                dictateBtn.textContent = '⌨️ Ej tillgänglig';
                dictateBtn.disabled = true;
                return;
            }}
            const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
            recognition = new SR();
            recognition.lang = 'sv-SE';
            recognition.continuous = true;
            recognition.interimResults = true;

            recognition.onstart = () => {{
                isRecording = true;
                dictateBtn.textContent = '⏹️ Stoppa';
                dictateBtn.classList.add('recording');
            }};

            recognition.onend = () => {{
                isRecording = false;
                dictateBtn.textContent = '🎤 Diktera';
                dictateBtn.classList.remove('recording');
                const text = textArea.value.trim();
                if (text && !text.endsWith('.') && !text.endsWith('!') && !text.endsWith('?') && !text.endsWith(':') && !text.endsWith(']')) {{
                    textArea.value = text + '.';
                }}
                saveLocal();
            }};

            recognition.onresult = (e) => {{
                for (let i = e.resultIndex; i < e.results.length; i++) {{
                    if (e.results[i].isFinal) {{
                        if (textArea.value && !textArea.value.endsWith(' ')) textArea.value += ' ';
                        textArea.value += e.results[i][0].transcript;
                        saveLocal();
                    }}
                }}
                updateWordCount();
            }};

            recognition.onerror = (e) => {{
                if (e.error === 'not-allowed') showToast('Ge mikrofontillåtelse');
                isRecording = false;
                dictateBtn.textContent = '🎤 Diktera';
                dictateBtn.classList.remove('recording');
            }};
        }}

        function toggleDictation() {{
            if (!recognition) {{ showToast('Diktering ej tillgänglig'); return; }}
            isRecording ? recognition.stop() : recognition.start();
        }}

        let photoInputBusy = false;
        let imageCount = 0;  // Räknare för bilder (lagras på server)

        function compressImage(file, callback) {{
            const img = new Image();
            img.onload = function() {{
                const canvas = document.createElement('canvas');
                const MAX = 1024;
                let w = img.width, h = img.height;
                if (w > h && w > MAX) {{ h = h * MAX / w; w = MAX; }}
                else if (h > MAX) {{ w = w * MAX / h; h = MAX; }}
                canvas.width = w;
                canvas.height = h;
                canvas.getContext('2d').drawImage(img, 0, 0, w, h);
                const compressed = canvas.toDataURL('image/jpeg', 0.5);
                URL.revokeObjectURL(img.src);
                callback(compressed);
            }};
            img.src = URL.createObjectURL(file);
        }}

        function takePhoto() {{
            if (photoInputBusy) {{
                showToast('Vänta...');
                return;
            }}
            photoInputBusy = true;

            const newInput = document.createElement('input');
            newInput.type = 'file';
            newInput.accept = 'image/*';
            newInput.capture = 'environment';
            newInput.style.display = 'none';
            newInput.onchange = function(e) {{
                const file = e.target.files[0];
                if (!file) {{
                    newInput.remove();
                    photoInputBusy = false;
                    return;
                }}
                compressImage(file, async function(compressedData) {{
                    try {{
                        // Skicka bilden direkt till servern
                        const resp = await fetch(IMAGE_URL, {{
                            method: 'POST',
                            headers: {{ 'Content-Type': 'application/json' }},
                            body: JSON.stringify({{
                                username: USERNAME,
                                session_id: getSessionId(),
                                image: compressedData
                            }})
                        }});
                        const data = await resp.json();
                        if (data.status === 'ok') {{
                            imageCount = data.image_count;
                            const num = imageCount;
                            if (textArea.value && !textArea.value.endsWith(' ')) textArea.value += ' ';
                            textArea.value += '[BILD' + num + ']';
                            saveLocal();
                            renderImageCounter();
                            showToast('Bild ' + num + ' uppladdad');
                        }} else {{
                            showToast('Fel: ' + (data.error || 'okänt'));
                        }}
                    }} catch (err) {{
                        showToast('Uppladdning misslyckades');
                    }}
                    newInput.remove();
                    photoInputBusy = false;
                }});
            }};
            document.body.appendChild(newInput);
            newInput.click();
            // Återställ om användaren avbryter (ingen change-event inom 60 sek)
            setTimeout(() => {{
                if (photoInputBusy) {{
                    photoInputBusy = false;
                    newInput.remove();
                }}
            }}, 60000);
        }}

        function renderImageCounter() {{
            imagePreview.innerHTML = '';
            if (imageCount === 0) return;
            const counter = document.createElement('div');
            counter.className = 'image-counter-big';
            counter.textContent = imageCount + ' bilder';
            imagePreview.appendChild(counter);
            updateWordCount();
        }}

        async function saveToWord() {{
            const text = textArea.value.trim();
            if (!text && imageCount === 0) {{ showToast('Ingen text att spara'); return; }}
            if (!navigator.onLine) {{ showToast('Kräver internet'); return; }}

            saveBtn.disabled = true;
            saveBtn.textContent = '⏳ Sparar...';

            try {{
                const sid = getSessionId();

                if (text) {{
                    await fetch(ACCUMULATE_URL, {{
                        method: 'POST',
                        headers: {{ 'Content-Type': 'application/json' }},
                        body: JSON.stringify({{ username: USERNAME, session_id: sid, text }})
                    }});
                }}

                // Bilder är redan uppladdade till servern
                saveBtn.textContent = '⏳ Skapar Word...';
                const resp = await fetch(WORD_URL, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ username: USERNAME, session_id: sid }})
                }});

                if (resp.ok) {{
                    const blob = await resp.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = USERNAME + '-' + new Date().toISOString().slice(0,10) + '.docx';
                    a.click();
                    URL.revokeObjectURL(url);
                    showToast('Word-fil nedladdad!');
                    localStorage.removeItem(SESSION_KEY);
                    getSessionId();
                }} else {{
                    throw new Error('Kunde inte skapa Word');
                }}
            }} catch (err) {{
                showToast('Fel: ' + err.message);
            }} finally {{
                saveBtn.disabled = false;
                saveBtn.textContent = '💾 Spara till Word';
            }}
        }}

        async function clearText() {{
            if (confirm('Rensa allt?')) {{
                textArea.value = 'Objekt: ';
                imageCount = 0;
                localStorage.removeItem(STORAGE_KEY);
                renderImageCounter();
                updateWordCount();
                // Rensa även på servern
                try {{
                    await fetch(IMAGE_URL, {{
                        method: 'POST',
                        headers: {{ 'Content-Type': 'application/json' }},
                        body: JSON.stringify({{
                            username: USERNAME,
                            session_id: getSessionId(),
                            action: 'clear'
                        }})
                    }});
                }} catch (e) {{
                    console.error('Kunde inte rensa server:', e);
                }}
                showToast('Rensat');
            }}
        }}

        // PIN-logik
        const PIN_KEY = 'selatek-' + USERNAME + '-pin';
        const pinOverlay = document.getElementById('pinOverlay');
        const pinInput = document.getElementById('pinInput');
        const pinTitle = document.getElementById('pinTitle');
        const pinError = document.getElementById('pinError');
        const pinInfo = document.getElementById('pinInfo');
        const pinBtn = document.getElementById('pinBtn');
        let pinMode = 'verify'; // 'verify', 'create', 'change'
        let newPin = '';

        function getLocalPin() {{
            return localStorage.getItem(PIN_KEY);
        }}

        function setLocalPin(pin) {{
            localStorage.setItem(PIN_KEY, pin);
        }}

        function showPinOverlay(mode) {{
            pinMode = mode;
            pinError.textContent = '';
            pinInput.value = '';
            newPin = '';

            if (mode === 'verify') {{
                pinTitle.textContent = 'Ange PIN-kod';
                pinInfo.textContent = '4 siffror';
            }} else if (mode === 'create') {{
                pinTitle.textContent = 'Skapa PIN-kod';
                pinInfo.textContent = 'Välj 4 siffror';
            }} else if (mode === 'change') {{
                pinTitle.textContent = 'Ange nuvarande PIN';
                pinInfo.textContent = '';
            }} else if (mode === 'new') {{
                pinTitle.textContent = 'Ny PIN-kod';
                pinInfo.textContent = 'Välj 4 siffror';
            }} else if (mode === 'confirm') {{
                pinTitle.textContent = 'Bekräfta PIN-kod';
                pinInfo.textContent = 'Skriv igen';
            }}

            pinOverlay.classList.remove('hidden');
            pinInput.focus();
        }}

        function hidePinOverlay() {{
            pinOverlay.classList.add('hidden');
        }}

        function submitPin() {{
            const pin = pinInput.value;

            if (pin.length !== 4 || !/^\d+$/.test(pin)) {{
                pinError.textContent = 'Ange 4 siffror';
                return;
            }}

            if (pinMode === 'verify') {{
                // Verifiera mot server om vi inte har lokal PIN
                const localPin = getLocalPin();
                if (localPin && pin === localPin) {{
                    hidePinOverlay();
                    showToast('Välkommen!');
                }} else if (HAS_SERVER_PIN) {{
                    // Verifiera mot servern
                    verifyPinWithServer(pin);
                }} else {{
                    pinError.textContent = 'Fel PIN-kod';
                    pinInput.value = '';
                }}
            }} else if (pinMode === 'create' || pinMode === 'new') {{
                newPin = pin;
                pinInput.value = '';
                pinMode = 'confirm';
                pinTitle.textContent = 'Bekräfta PIN-kod';
                pinInfo.textContent = 'Skriv igen';
            }} else if (pinMode === 'confirm') {{
                if (pin === newPin) {{
                    setLocalPin(pin);
                    saveServerPin(pin);
                    hidePinOverlay();
                    showToast('PIN-kod sparad!');
                    pinBtn.textContent = '🔐 Ändra PIN-kod';
                }} else {{
                    pinError.textContent = 'PIN matchar inte';
                    pinInput.value = '';
                    newPin = '';
                    pinMode = 'create';
                    pinTitle.textContent = 'Skapa PIN-kod';
                    pinInfo.textContent = 'Försök igen';
                }}
            }} else if (pinMode === 'change') {{
                // Verifiera nuvarande PIN mot servern
                verifyPinForChange(pin);
            }}
        }}

        async function saveServerPin(pin) {{
            try {{
                await fetch(ACCUMULATE_URL, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ username: USERNAME, action: 'set_pin', pin: pin }})
                }});
            }} catch (e) {{
                console.error('Kunde inte spara PIN på server:', e);
            }}
        }}

        async function verifyPinWithServer(pin) {{
            try {{
                const resp = await fetch(ACCUMULATE_URL, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ username: USERNAME, action: 'verify_pin', pin: pin }})
                }});
                const data = await resp.json();
                if (data.valid) {{
                    // PIN korrekt - spara lokalt och fortsätt
                    setLocalPin(pin);
                    hidePinOverlay();
                    showToast('Välkommen!');
                }} else {{
                    pinError.textContent = 'Fel PIN-kod';
                    pinInput.value = '';
                }}
            }} catch (e) {{
                pinError.textContent = 'Kunde inte verifiera';
                pinInput.value = '';
            }}
        }}

        async function verifyPinForChange(pin) {{
            try {{
                const resp = await fetch(ACCUMULATE_URL, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ username: USERNAME, action: 'verify_pin', pin: pin }})
                }});
                const data = await resp.json();
                if (data.valid) {{
                    // Nuvarande PIN korrekt - gå vidare till ny PIN
                    pinInput.value = '';
                    pinMode = 'new';
                    pinTitle.textContent = 'Ny PIN-kod';
                    pinInfo.textContent = 'Välj 4 siffror';
                    pinError.textContent = '';
                }} else {{
                    pinError.textContent = 'Fel PIN-kod';
                    pinInput.value = '';
                }}
            }} catch (e) {{
                pinError.textContent = 'Kunde inte verifiera';
                pinInput.value = '';
            }}
        }}

        function handlePinBtn() {{
            const localPin = getLocalPin();
            if (localPin) {{
                showPinOverlay('change');
            }} else {{
                showPinOverlay('create');
            }}
        }}

        function initPin() {{
            const localPin = getLocalPin();

            if (HAS_SERVER_PIN && !localPin) {{
                // Server har PIN men vi har ingen lokalt - be om PIN
                showPinOverlay('verify');
                pinInfo.textContent = 'Du har en PIN på servern. Ange den eller kontakta admin.';
            }} else if (localPin) {{
                // Vi har lokal PIN - verifiera
                showPinOverlay('verify');
            }} else {{
                // Ingen PIN - göm overlay
                hidePinOverlay();
            }}

            pinBtn.textContent = localPin || HAS_SERVER_PIN ? '🔐 Ändra PIN-kod' : '🔐 Skapa PIN-kod';
        }}

        // Enter-tangent för PIN-input
        pinInput.addEventListener('keypress', (e) => {{
            if (e.key === 'Enter') submitPin();
        }});

        pinBtn.addEventListener('click', handlePinBtn);

        dictateBtn.addEventListener('click', toggleDictation);
        photoBtn.addEventListener('click', takePhoto);
        saveBtn.addEventListener('click', saveToWord);
        clearBtn.addEventListener('click', clearText);
        textArea.addEventListener('input', saveLocal);
        window.addEventListener('online', updateStatus);
        window.addEventListener('offline', updateStatus);

        async function loadImageCountFromServer() {{
            const sid = getSessionId();
            console.log('Hämtar bildräkning för session:', sid);
            try {{
                const resp = await fetch(IMAGE_URL, {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{
                        username: USERNAME,
                        session_id: sid,
                        action: 'get_count'
                    }})
                }});
                const data = await resp.json();
                console.log('Bildräkning från server:', data);
                if (data.status === 'ok') {{
                    imageCount = data.image_count;
                    renderImageCounter();
                }}
            }} catch (e) {{
                console.error('Kunde inte hämta bildräkning:', e);
            }}
        }}

        loadLocal();
        updateStatus();
        updateWordCount();
        setupSpeechRecognition();
        initPin();
        loadImageCountFromServer();
    </script>
</body>
</html>'''


def get_admin_html():
    """Admin-sida"""
    return '''<!DOCTYPE html>
<html lang="sv">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Selatek Admin</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            min-height: 100vh; color: #fff; padding: 20px;
        }
        .container { max-width: 500px; margin: 0 auto; }
        h1 { text-align: center; margin-bottom: 30px; }
        .card { background: rgba(255,255,255,0.1); border-radius: 12px; padding: 20px; margin-bottom: 20px; }
        .card h2 { margin-bottom: 15px; font-size: 18px; }
        input {
            width: 100%; padding: 12px;
            border: 1px solid rgba(255,255,255,0.2); border-radius: 8px;
            background: rgba(255,255,255,0.1); color: #fff;
            font-size: 16px; margin-bottom: 10px;
        }
        input::placeholder { color: rgba(255,255,255,0.5); }
        button {
            width: 100%; padding: 14px; border: none; border-radius: 8px;
            background: linear-gradient(135deg, #10b981 0%, #047857 100%);
            color: #fff; font-size: 16px; font-weight: 600; cursor: pointer;
        }
        button:hover { opacity: 0.9; }
        .user-list { list-style: none; }
        .user-list li {
            display: flex; justify-content: space-between; align-items: center;
            padding: 12px; background: rgba(255,255,255,0.05);
            border-radius: 8px; margin-bottom: 8px;
        }
        .user-list a { color: #60a5fa; text-decoration: none; }
        .user-list .delete-btn {
            background: #ef4444; border: none; color: #fff;
            padding: 6px 12px; border-radius: 6px; cursor: pointer; font-size: 12px;
        }
        .message { padding: 12px; border-radius: 8px; margin-bottom: 15px; text-align: center; }
        .message.success { background: rgba(16, 185, 129, 0.2); color: #4ade80; }
        .message.error { background: rgba(239, 68, 68, 0.2); color: #f87171; }
        .no-users { text-align: center; color: #888; padding: 20px; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🔐 Selatek Admin</h1>
        <div class="card">
            <h2>Skapa ny användare</h2>
            <div id="message"></div>
            <input type="password" id="adminSecret" placeholder="Admin-lösenord">
            <input type="text" id="newUsername" placeholder="Användarnamn (t.ex. marcus)">
            <button onclick="createUser()">Skapa användare</button>
        </div>
        <div class="card">
            <h2>Aktiva användare</h2>
            <ul class="user-list" id="userList"><li class="no-users">Laddar...</li></ul>
        </div>
    </div>
    <script>
        const USERS_URL = 'https://mackanh1972--selatek-notes-multiuser-list-users.modal.run';
        const CREATE_URL = 'https://mackanh1972--selatek-notes-multiuser-create-user.modal.run';
        const DELETE_URL = 'https://mackanh1972--selatek-notes-multiuser-delete-user.modal.run';
        const USER_BASE = 'https://mackanh1972--selatek-notes-multiuser-user-page.modal.run';

        async function loadUsers() {
            try {
                const resp = await fetch(USERS_URL);
                const data = await resp.json();
                const list = document.getElementById('userList');
                if (data.users && data.users.length > 0) {
                    list.innerHTML = data.users.map(u => `
                        <li>
                            <a href="${USER_BASE}?username=${u}" target="_blank">${u}</a>
                            <button class="delete-btn" onclick="deleteUser('${u}')">Ta bort</button>
                        </li>
                    `).join('');
                } else {
                    list.innerHTML = '<li class="no-users">Inga användare ännu</li>';
                }
            } catch (e) { console.error(e); }
        }

        async function createUser() {
            const secret = document.getElementById('adminSecret').value;
            const username = document.getElementById('newUsername').value.toLowerCase().trim();
            const msg = document.getElementById('message');

            if (!secret || !username) {
                msg.className = 'message error';
                msg.textContent = 'Fyll i båda fälten';
                return;
            }
            if (!/^[a-z0-9-]+$/.test(username)) {
                msg.className = 'message error';
                msg.textContent = 'Endast a-z, 0-9, bindestreck';
                return;
            }

            try {
                const resp = await fetch(CREATE_URL, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ secret, username })
                });
                const data = await resp.json();
                if (data.status === 'ok') {
                    msg.className = 'message success';
                    msg.textContent = 'Skapad: ' + username;
                    document.getElementById('newUsername').value = '';
                    loadUsers();
                } else {
                    msg.className = 'message error';
                    msg.textContent = data.error || 'Fel';
                }
            } catch (e) {
                msg.className = 'message error';
                msg.textContent = e.message;
            }
        }

        async function deleteUser(username) {
            const secret = document.getElementById('adminSecret').value;
            if (!secret) { alert('Ange lösenord först'); return; }
            if (!confirm('Ta bort ' + username + '?')) return;

            try {
                const resp = await fetch(DELETE_URL, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ secret, username })
                });
                const data = await resp.json();
                if (data.status === 'ok') loadUsers();
                else alert(data.error || 'Fel');
            } catch (e) { alert(e.message); }
        }

        loadUsers();
    </script>
</body>
</html>'''


# ============ ENDPOINTS ============

@app.function(image=image)
@modal.fastapi_endpoint(method="GET")
def admin():
    """Admin-sida"""
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=get_admin_html())


@app.function(image=image)
@modal.fastapi_endpoint(method="GET")
def list_users():
    """Lista användare"""
    try:
        user_list = users.get("_user_list", [])
        return {"users": user_list}
    except:
        return {"users": []}


@app.function(image=image)
@modal.fastapi_endpoint(method="POST")
def create_user(data: dict):
    """Skapa användare"""
    secret = data.get("secret", "")
    username = data.get("username", "").lower().strip()

    if secret != ADMIN_SECRET:
        return {"status": "error", "error": "Fel lösenord"}

    if not username or len(username) < 2:
        return {"status": "error", "error": "Minst 2 tecken"}

    user_list = users.get("_user_list", [])
    if username in user_list:
        return {"status": "error", "error": "Finns redan"}

    user_list.append(username)
    users["_user_list"] = user_list
    users[username] = {"created": str(datetime.now()), "pin": "9999"}  # Standard-PIN

    return {"status": "ok", "username": username}


@app.function(image=image)
@modal.fastapi_endpoint(method="POST")
def delete_user(data: dict):
    """Ta bort användare"""
    secret = data.get("secret", "")
    username = data.get("username", "").lower().strip()

    if secret != ADMIN_SECRET:
        return {"status": "error", "error": "Fel lösenord"}

    user_list = users.get("_user_list", [])
    if username not in user_list:
        return {"status": "error", "error": "Finns inte"}

    user_list.remove(username)
    users["_user_list"] = user_list
    try:
        del users[username]
    except:
        pass

    return {"status": "ok"}


@app.function(image=image)
@modal.fastapi_endpoint(method="GET")
def user_page(username: str = ""):
    """Användarsida"""
    from fastapi.responses import HTMLResponse, JSONResponse

    username = username.lower().strip()
    if not username:
        return JSONResponse(status_code=400, content={"error": "Ange username"})

    user_list = users.get("_user_list", [])
    if username not in user_list:
        return JSONResponse(status_code=404, content={"error": "Användaren finns inte"})

    # Kolla om användaren har PIN
    user_data = users.get(username, {})
    has_pin = bool(user_data.get("pin"))

    return HTMLResponse(content=get_user_html(username, has_pin))


@app.function(image=image)
@modal.fastapi_endpoint(method="POST")
def accumulate(data: dict):
    """Spara text eller sätt PIN"""
    from fastapi.responses import JSONResponse

    username = data.get("username", "").lower().strip()
    action = data.get("action", "")

    user_list = users.get("_user_list", [])
    if username not in user_list:
        return JSONResponse(status_code=404, content={"error": "Användaren finns inte"})

    # Hantera PIN-actions
    if action == "set_pin":
        pin = data.get("pin", "")
        if len(pin) == 4 and pin.isdigit():
            user_data = users.get(username, {"created": str(datetime.now())})
            user_data["pin"] = pin
            users[username] = user_data
            return {"status": "ok", "message": "PIN sparad"}
        return {"status": "error", "error": "Ogiltig PIN"}

    if action == "verify_pin":
        pin = data.get("pin", "")
        user_data = users.get(username, {})
        server_pin = user_data.get("pin", "9999")  # Standard-PIN om ingen finns
        if pin == server_pin:
            return {"status": "ok", "valid": True}
        return {"status": "ok", "valid": False}

    # Vanlig text-lagring - spara i users dict istället för sessions
    text = data.get("text", "")

    user_data = users.get(username, {})
    user_data["text"] = text
    users[username] = user_data

    return {"status": "ok", "text_length": len(text)}


@app.function(image=image)
@modal.fastapi_endpoint(method="POST")
def add_image(data: dict):
    """Lägg till bild eller hämta bildräkning"""
    from fastapi.responses import JSONResponse

    username = data.get("username", "").lower().strip()
    action = data.get("action", "")
    image_data = data.get("image", "")

    user_list = users.get("_user_list", [])
    if username not in user_list:
        return JSONResponse(status_code=404, content={"error": "Användaren finns inte"})

    # Använd users dict för att spara bilder (mer persistent)
    user_data = users.get(username, {})
    user_images = user_data.get("images", [])

    # Hämta bara bildräkning utan att lägga till bild
    if action == "get_count":
        return {"status": "ok", "image_count": len(user_images)}

    # Rensa bilderna
    if action == "clear":
        user_data["images"] = []
        user_data["text"] = ""
        users[username] = user_data
        return {"status": "ok", "message": "Session rensad"}

    # Lägg till bild
    user_images.append(image_data)
    user_data["images"] = user_images
    users[username] = user_data

    return {"status": "ok", "image_count": len(user_images)}


@app.function(image=image)
@modal.fastapi_endpoint(method="POST")
def save_to_word(data: dict):
    """Skapa Word-dokument"""
    import base64
    import io
    import re
    from fastapi.responses import Response, JSONResponse
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from PIL import Image

    username = data.get("username", "").lower().strip()

    user_list = users.get("_user_list", [])
    if username not in user_list:
        return JSONResponse(status_code=404, content={"error": "Användaren finns inte"})

    # Hämta från users dict istället för sessions
    user_data = users.get(username, {})
    text = user_data.get("text", "")
    images_data = user_data.get("images", [])

    doc = Document()

    def remove_indent(p):
        pPr = p._p.get_or_add_pPr()
        existing = pPr.find(qn('w:ind'))
        if existing is not None:
            pPr.remove(existing)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:right'), '0')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:hanging'), '0')
        pPr.append(ind)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_indent(title)
    title_run = title.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M")}')
    title_run.bold = True
    title_run.font.size = Pt(16)

    pattern = r'(\[BILD\d+\])'
    parts = re.split(pattern, text)

    for part in parts:
        part = part.strip()
        if not part:
            continue

        match = re.match(r'\[BILD(\d+)\]', part)
        if match:
            img_num = int(match.group(1))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True  # Håll ihop bildtext med bilden på samma sida
            remove_indent(p)
            run = p.add_run(f'BILD {img_num}')
            run.bold = True
            run.font.size = Pt(11)

            if img_num <= len(images_data):
                try:
                    img_d = images_data[img_num - 1]
                    if ',' in img_d:
                        img_d = img_d.split(',')[1]

                    img_bytes = base64.b64decode(img_d)
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    rotated = pil_img.rotate(0, expand=True)  # Ingen rotation

                    buf = io.BytesIO()
                    rotated.save(buf, format='JPEG', quality=85)
                    buf.seek(0)

                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    remove_indent(pic_p)
                    pic_p.add_run().add_picture(buf, width=Inches(3))
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f'[Fel bild {img_num}: {e}]')
        else:
            p = doc.add_paragraph(part)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            remove_indent(p)

    doc_buf = io.BytesIO()
    doc.save(doc_buf)
    doc_buf.seek(0)

    # Rensa inte automatiskt - låt användaren göra det manuellt med Rensa-knappen

    return Response(
        content=doc_buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'inline; filename="{username}-{datetime.now().strftime("%Y-%m-%d")}.docx"'}
    )


# Health endpoint borttagen för att hålla oss under 8 endpoint-gränsen
