"""
exp-002: Session-baserad textackumulering med Word-export och bilder
Hypothesis: Modal Volume + session ID kan behålla state mellan anrop
"""
import modal
from pathlib import Path
from datetime import datetime
import base64
import json

# Skapa Volume för sessions
volume = modal.Volume.from_name("voice-sessions", create_if_missing=True)
image = modal.Image.debian_slim().pip_install("fastapi", "python-docx")
app = modal.App("voice-accumulator", image=image)

SESSIONS_DIR = Path("/sessions")


@app.function(volumes={"/sessions": volume})
@modal.fastapi_endpoint(method="POST", docs=True)
def accumulate(data: dict):
    """
    Ackumulera text för given session.

    Body: {"session_id": "...", "text": "..."}
    Returns: Plain text (all accumulated text) - ingen JSON, inga citattecken
    """
    from fastapi.responses import PlainTextResponse

    session_id = data.get("session_id", "")
    new_text = data.get("text", "").strip()

    if not session_id:
        return PlainTextResponse("Fel: session_id saknas", media_type="text/plain; charset=utf-8")

    # Skapa session-mapp om den inte finns
    session_dir = SESSIONS_DIR / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    session_file = session_dir / "text.txt"

    if not new_text:
        # Om ingen ny text, returnera bara existerande
        if session_file.exists():
            return PlainTextResponse(session_file.read_text(encoding="utf-8"), media_type="text/plain; charset=utf-8")
        return PlainTextResponse("", media_type="text/plain; charset=utf-8")

    # Läs existerande text
    existing = ""
    if session_file.exists():
        existing = session_file.read_text(encoding="utf-8")

    # Append ny text (med mellanslag om det finns befintlig text)
    if existing:
        accumulated = existing + " " + new_text
    else:
        accumulated = new_text

    # Skriv tillbaka
    session_file.write_text(accumulated, encoding="utf-8")
    volume.commit()

    # Returnera plain text utan citattecken, med UTF-8 för åäö
    return PlainTextResponse(accumulated, media_type="text/plain; charset=utf-8")


@app.function(volumes={"/sessions": volume})
@modal.fastapi_endpoint(method="POST", docs=True)
def add_image(data: dict):
    """
    Lägg till en bild till sessionen.

    Body: {"session_id": "...", "image": "base64-encoded-image"}
    Returns: Bildnummer och uppdaterad text
    """
    from fastapi.responses import JSONResponse

    session_id = data.get("session_id", "")
    image_data = data.get("image", "")

    if not session_id:
        return JSONResponse({"error": "session_id saknas"}, status_code=400)

    if not image_data:
        return JSONResponse({"error": "image saknas"}, status_code=400)

    # Skapa session-mapp
    session_dir = SESSIONS_DIR / session_id
    session_dir.mkdir(parents=True, exist_ok=True)

    # Räkna befintliga bilder
    existing_images = list(session_dir.glob("bild_*.jpg"))
    image_num = len(existing_images) + 1

    # Spara bilden
    image_filename = f"bild_{image_num}.jpg"
    image_path = session_dir / image_filename

    # Avkoda base64 och spara
    try:
        # Ta bort data:image/jpeg;base64, prefix om det finns
        if "," in image_data:
            image_data = image_data.split(",")[1]
        image_bytes = base64.b64decode(image_data)
        image_path.write_bytes(image_bytes)
    except Exception as e:
        return JSONResponse({"error": f"Kunde inte spara bild: {str(e)}"}, status_code=400)

    # Lägg till bildreferens i texten
    session_file = session_dir / "text.txt"
    existing = ""
    if session_file.exists():
        existing = session_file.read_text(encoding="utf-8")

    # Lägg till [BILD X] i texten
    if existing:
        updated = existing + f" [BILD{image_num}]"
    else:
        updated = f"[BILD{image_num}]"

    session_file.write_text(updated, encoding="utf-8")
    volume.commit()

    return JSONResponse({
        "success": True,
        "image_num": image_num,
        "text": updated
    })


@app.function(volumes={"/sessions": volume})
@modal.fastapi_endpoint(method="POST", docs=True)
def save_to_word(data: dict):
    """
    Spara session till Word-dokument med bilder.

    Body: {"session_id": "..."}
    Returns: Word-dokument som bytes
    """
    from docx import Document
    from docx.shared import Inches
    from fastapi.responses import Response
    from io import BytesIO

    session_id = data.get("session_id", "")

    if not session_id:
        return Response(content="Fel: session_id saknas", media_type="text/plain")

    session_dir = SESSIONS_DIR / session_id
    session_file = session_dir / "text.txt"

    if not session_file.exists():
        return Response(content="Fel: Ingen text sparad", media_type="text/plain")

    text = session_file.read_text(encoding="utf-8")

    # Skapa Word-dokument
    doc = Document()
    doc.add_heading("Dikterad text", 0)
    doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # Lägg till texten
    doc.add_paragraph(text)

    # Hitta alla bilder
    images = sorted(session_dir.glob("bild_*.jpg"))

    if images:
        doc.add_paragraph("")
        doc.add_heading("Bilagor", level=1)

        for img_path in images:
            # Extrahera bildnummer från filnamn (bild_1.jpg -> 1)
            img_num = img_path.stem.split("_")[1]

            # Skapa bildtext-stycke med keep_with_next för att hålla ihop med bilden
            label_para = doc.add_paragraph(f"BILD{img_num}:")
            label_para.paragraph_format.keep_with_next = True  # Håll ihop med nästa element (bilden)

            try:
                doc.add_picture(str(img_path), width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"[Kunde inte ladda bild: {e}]")

            doc.add_paragraph("")

    # Spara till bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Rensa sessionen efter export
    try:
        for f in session_dir.iterdir():
            f.unlink()
        session_dir.rmdir()
        volume.commit()
    except Exception:
        pass

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=anteckning_{session_id[:8]}.docx"}
    )


@app.function()
@modal.fastapi_endpoint(method="GET", docs=True)
def health() -> dict:
    return {"status": "ok", "experiment": "exp-002", "features": ["text", "images", "word-export"]}
